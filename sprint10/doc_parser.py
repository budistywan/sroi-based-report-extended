"""
doc_parser.py — Sprint 10A
Parser untuk dokumen .docx SROI menjadi parsed_source_json.

Preserves:
  - Heading hierarchy sebagai heading_path
  - Paragraf per blok
  - Tabel sebagai struktur tabel
  - Semantic hints heuristik (tidak overclaim)

Usage:
  python doc_parser.py --input /path/file.docx --output /path/parsed.json
  INPUT_FILE=... OUTPUT_FILE=... python doc_parser.py
"""

import json, sys, os, re, argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    print("FAIL: python-docx tidak terinstall — jalankan: pip install python-docx")
    sys.exit(1)

PARSER_VERSION = "1.0.0"

# ── PATH CONFIG ───────────────────────────────────────────────
parser = argparse.ArgumentParser()
parser.add_argument("--input",  default=None)
parser.add_argument("--output", default=None)
args = parser.parse_args()

SCRIPT_DIR  = Path(__file__).parent
INPUT_FILE  = Path(args.input)  if args.input  \
    else Path(os.environ.get("INPUT_FILE",
              "/mnt/user-data/uploads/1775796870673_17122025_SROI_Kresna_Patra_FT_Boyolali__1_.docx"))
OUTPUT_FILE = Path(args.output) if args.output \
    else Path(os.environ.get("OUTPUT_FILE", SCRIPT_DIR / "parsed_kresna_doc.json"))

print(f"Input  : {INPUT_FILE.resolve()}")
print(f"Output : {OUTPUT_FILE.resolve()}")

if not INPUT_FILE.exists():
    print(f"FAIL: {INPUT_FILE} tidak ditemukan"); sys.exit(1)

# ── SEMANTIC HINT RULES ───────────────────────────────────────
# Heuristik — keyword-based, tidak authoritative
SEMANTIC_KEYWORDS = {
    "baseline_or_context": [
        "kondisi awal", "baseline", "konteks", "latar belakang",
        "gambaran umum", "kondisi saat ini", "existing",
    ],
    "potential": [
        "potensi", "peluang", "potential", "prospek",
        "kekuatan", "keunggulan",
    ],
    "problem_framing": [
        "masalah", "hambatan", "barrier", "permasalahan",
        "tantangan", "kendala", "kelemahan",
    ],
    "strategy_roadmap": [
        "strategi", "roadmap", "rencana", "tahapan", "langkah",
        "pendekatan", "metode", "intervensi",
    ],
    "lfa": [
        "lfa", "logical framework", "logframe", "kerangka logis",
        "rantai hasil", "theory of change",
    ],
    "toc_pathway": [
        "jalur perubahan", "theory of change", "toc",
        "pathway", "rantai kausal", "causal chain",
        "mekanisme perubahan",
    ],
    "ddat_justification": [
        "deadweight", "displacement", "attribution", "drop-off",
        "ddat", "fiksasi", "haircut", "adjustment",
    ],
    "monetization": [
        "monetisasi", "nilai sosial", "nilai ekonomi",
        "proxy", "financial proxy", "wtp", "willingness to pay",
    ],
    "sroi_result": [
        "sroi", "social return", "rasio", "ratio",
        "return on investment", "nilai investasi",
    ],
    "learning_recommendation": [
        "pembelajaran", "rekomendasi", "lesson learned",
        "temuan", "kesimpulan", "penutup", "saran",
    ],
}

def get_semantic_hint(text: str) -> str | None:
    """Kembalikan semantic hint berdasarkan keyword matching.
       Ambiguous → None. Tidak overclaim."""
    text_lower = text.lower()
    matches = []
    for hint, keywords in SEMANTIC_KEYWORDS.items():
        if any(kw in text_lower for kw in keywords):
            matches.append(hint)
    if len(matches) == 1:
        return matches[0]
    if len(matches) > 1:
        return matches[0]  # ambil yang pertama match
    return None

def heading_level(style_name: str) -> int | None:
    """Ekstrak level dari heading style (Heading 1 → 1, dll)."""
    m = re.search(r'heading\s*(\d)', style_name, re.IGNORECASE)
    if m:
        return int(m.group(1))
    if style_name.lower().startswith('toc'):
        return None  # skip TOC entries
    return None

# ── LOAD DOCUMENT ─────────────────────────────────────────────
print("Loading document...")
doc = Document(str(INPUT_FILE))
print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# ── EXTRACT DOCUMENT TITLE ────────────────────────────────────
def extract_title(doc) -> str:
    """Cari judul dari metadata atau heading pertama."""
    props = doc.core_properties
    if props.title:
        return props.title.strip()
    # Fallback: heading 1 pertama yang bukan sampul
    for p in doc.paragraphs:
        if heading_level(p.style.name) == 1 and p.text.strip():
            txt = p.text.strip()
            if len(txt) > 5:
                return txt
    return ""

document_title = extract_title(doc)
print(f"  Title: {document_title[:60]}")

# ── EXTRACT TABLES ─────────────────────────────────────────────
def extract_table(tbl, table_idx: int) -> dict:
    rows = []
    for row in tbl.rows:
        cells = [c.text.strip() for c in row.cells]
        # Deduplicate merged cells
        deduped = []
        prev = None
        for c in cells:
            if c != prev:
                deduped.append(c)
            prev = c
        rows.append(deduped)

    # Coba deteksi header row
    header = rows[0] if rows else []
    data_rows = rows[1:] if len(rows) > 1 else []

    # Semantic hint dari konten tabel
    all_text = " ".join(" ".join(r) for r in rows[:3]).lower()
    hint = get_semantic_hint(all_text)

    return {
        "table_idx":    table_idx,
        "rows":         len(tbl.rows),
        "cols":         len(tbl.columns),
        "header":       header,
        "data_rows":    data_rows[:20],   # cap 20 rows untuk efisiensi
        "total_rows":   len(rows),
        "semantic_hint": hint,
    }

print("Extracting tables...")
tables_out = [extract_table(t, i) for i, t in enumerate(doc.tables)]
print(f"  {len(tables_out)} tables extracted")

# ── EXTRACT SECTIONS ──────────────────────────────────────────
# Track heading path sebagai stack
heading_stack: list[str] = []
sections_out  = []
para_index    = 0  # track posisi dalam dokumen
skip_toc      = False

print("Extracting sections...")
for para in doc.paragraphs:
    style = para.style.name
    text  = para.text.strip()

    # Skip empty
    if not text:
        continue

    level = heading_level(style)

    # Deteksi dan skip TOC
    if style.lower().startswith('toc') or text in ['DAFTAR ISI', 'Daftar Isi']:
        skip_toc = True
        continue
    if skip_toc and re.match(r'^(BAB|[IVX]+\.|[0-9]+\.)', text):
        skip_toc = False  # mulai konten setelah TOC

    if level is not None:
        # Update heading stack
        while len(heading_stack) >= level:
            heading_stack.pop() if heading_stack else None
        heading_stack.append(text)

        sections_out.append({
            "heading_path":  list(heading_stack),
            "block_type":    f"heading_{level}",
            "text":          text,
            "semantic_hint": get_semantic_hint(text),
        })
    else:
        # Paragraph atau list item
        hint = get_semantic_hint(text)
        if len(text) >= 10:  # skip teks yang terlalu pendek
            sections_out.append({
                "heading_path":  list(heading_stack),
                "block_type":    "paragraph",
                "text":          text,
                "semantic_hint": hint,
            })

print(f"  {len(sections_out)} sections extracted")

# ── DETECT DOCUMENT SIGNALS ───────────────────────────────────
all_text_lower = " ".join(
    s["text"].lower() for s in sections_out
)

def has_signal(keywords: list) -> bool:
    return any(kw in all_text_lower for kw in keywords)

signals = {
    "has_lfa":               has_signal(["lfa","logical framework","logframe","kerangka logis"]),
    "has_toc_narrative":     has_signal(["theory of change","jalur perubahan","toc","rantai kausal","causal chain"]),
    "has_ddat_justification":has_signal(["deadweight","displacement","attribution","drop-off","ddat","fiksasi dampak"]),
    "has_monetization":      has_signal(["monetisasi","financial proxy","proxy","nilai sosial","wtp"]),
    "has_sroi_calculation":  has_signal(["sroi","social return on investment","rasio sroi"]),
    "has_baseline":          has_signal(["kondisi awal","baseline","gambaran umum"]),
    "has_potential":         has_signal(["potensi","potential","peluang"]),
    "has_learning":          has_signal(["pembelajaran","rekomendasi","lesson learned","triple loop"]),
}
print("Signals:", {k: v for k, v in signals.items() if v})

# ── COMPOSE OUTPUT ────────────────────────────────────────────
output = {
    "source_type":    "docx",
    "document_title": document_title,
    "parsed_at":      datetime.now().isoformat(),
    "parser_version": PARSER_VERSION,
    "source_file":    INPUT_FILE.name,
    "stats": {
        "total_paragraphs": len(doc.paragraphs),
        "total_tables":     len(doc.tables),
        "sections_extracted": len(sections_out),
        "tables_extracted":   len(tables_out),
    },
    "signals":  signals,
    "sections": sections_out,
    "tables":   tables_out,
}

# ── WRITE ─────────────────────────────────────────────────────
OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
json.dump(output, open(OUTPUT_FILE,"w"), indent=2, ensure_ascii=False)

print(f"\nOutput: {OUTPUT_FILE}")
print(f"\n{'='*55}")
print("DOC PARSER COMPLETE")
print(f"  Sections : {len(sections_out)}")
print(f"  Tables   : {len(tables_out)}")
print(f"  Signals  : {sum(signals.values())} active")
print("="*55)
